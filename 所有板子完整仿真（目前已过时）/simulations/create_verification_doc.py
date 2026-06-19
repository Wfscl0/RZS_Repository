from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "五块板_LTspice_仿真验证报告.docx"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_font(paragraph, size=10.5, color=(45, 45, 45)):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=(45, 45, 45), align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=9.2, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_run_font(run, size=18 if level == 1 else 13, bold=True, color=(31, 78, 121))
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=9.8, color=(45, 45, 45))
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        shade_cell(hdr[i], "D9EAF7")
        set_cell_text(hdr[i], header, bold=True, color=(31, 78, 121), align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or value in ("PASS", "条件通过", "通过", "不通过") else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], str(value), align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    for row in table.rows:
        for cell in row.cells:
            cell.margin_left = Cm(0.08)
            cell.margin_right = Cm(0.08)
            cell.margin_top = Cm(0.05)
            cell.margin_bottom = Cm(0.05)
    return table


def add_figure(doc, image_name, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(ROOT / image_name), width=Inches(8.35))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=8.5, color=(90, 90, 90))
    cap.paragraph_format.space_after = Pt(2)


def add_board_section(doc, board):
    doc.add_page_break()
    add_heading(doc, board["title"], 1)
    add_body(doc, board["purpose"])
    add_table(doc, ["测试时段/测量点", "输入状态", "期望输出"], board["timing"], widths=[4.0, 9.0, 6.0])
    doc.add_paragraph()
    add_table(doc, ["测量项", "仿真值", "判定"], board["results"], widths=[5.3, 8.4, 5.3])
    add_body(doc, "结论：" + board["conclusion"])
    add_figure(doc, board["image"], board["caption"])
    for image_name, caption in board.get("extra_figures", []):
        add_figure(doc, image_name, caption)


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin = Cm(1.15)
section.bottom_margin = Cm(1.15)
section.left_margin = Cm(1.25)
section.right_margin = Cm(1.25)

styles = doc.styles
styles["Normal"].font.name = "Microsoft YaHei"
styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
styles["Normal"].font.size = Pt(9.8)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("五块板 LTspice 仿真验证报告")
set_run_font(run, size=24, bold=True, color=(31, 78, 121))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("依据《硬件电路接口.xlsx》、最新网表与原理图图片生成")
set_run_font(run, size=12, color=(90, 90, 90))

meta_rows = [
    ("仿真工具", "LTspice 26.0.1 for Windows"),
    ("仿真输入", "fin/*.tel、fin/*.png、fin/硬件电路接口.xlsx"),
    ("仿真输出", "simulations/*.cir、*.log、*.raw、*_waveform.png"),
    ("生成时间", "2026-05-18"),
]
add_table(doc, ["项目", "内容"], meta_rows, widths=[4.2, 14.8])

add_heading(doc, "1. 总体结论", 1)
add_body(doc, "本报告汇总五块板的接口时序仿真。每块板均以最新网表和原理图连接关系建立 LTspice 测试平台，使用官方模型或等效行为模型验证输入时序到接口输出的逻辑关系。")
summary_rows = [
    ("继电器检测", "PASS", "闭合时 *_ACTUAL 为 5 V、*# 为 0 V；断开后极性反转，符合接口表。"),
    ("EBS", "PASS", "正常输出 EBS_WORK=24 V；看门狗故障或 EBS_TRIG 故障时 EBS_WORK=0 V。"),
    ("锁存", "PASS", "IMD 故障后 SDC_OUT 拉低并保持锁存；复位低电平后恢复。"),
    ("BSPD", "条件通过", "故障条件成立后 314.29 ms 断开，小于 500 ms；冲突消失超过 10 s 后仍保持断开，BSPD 电源复位后恢复。"),
    ("TSAL", "PASS", "红灯按 2.775 Hz 闪烁，占空比 51.92%；安全状态下红灯关闭、绿灯点亮。"),
]
add_table(doc, ["板子", "结论", "摘要"], summary_rows, widths=[3.2, 2.6, 14.5])

add_heading(doc, "2. 说明与边界", 1)
for text in [
    "网表文件为封装/网络连接清单，原理图来源为 PNG 图片；本报告验证接口时序和逻辑极性，不等同于包含 PCB 寄生参数的完整电磁/模拟精度验证。",
    "未找到可直接用于 LTspice 的官方模型时，继电器触点、锁存状态、部分保护器件使用行为等效模型；这些等效模型仅用于验证接口逻辑。",
    "BSPD 的油压阈值在接口表中标注为“还没测定”，因此 BSPD 对“阈值已满足后的延时断开、保持断开和电源复位恢复行为”给出条件通过结论。",
]:
    add_body(doc, text)

boards = [
    {
        "title": "3. 继电器检测板",
        "purpose": "验证 MAINP、MAINN、PRE 三路继电器辅助接点检测的输出极性。三路拓扑一致，仿真使用一个代表通道：接点闭合时检测节点被上拉，断开后经 100 kΩ / 100 nF 放电。",
        "timing": [
            ("0-4 ms", "接点闭合，V(ctrl)=5 V", "检测节点高电平，*_ACTUAL=5 V，*#=0 V"),
            ("4.01 ms 起", "接点断开，V(ctrl)=0 V", "检测节点经 RC 放电，最终 *_ACTUAL=0 V，*#=5 V"),
            ("2 ms / 50 ms", "分别测闭合稳定值和断开稳定值", "用于判定接口表极性"),
        ],
        "results": [
            ("2 ms 闭合", "sense=4.733 V，ACTUAL=5.000 V，#=0 V", "PASS"),
            ("50 ms 断开", "sense=0.019 V，ACTUAL=0 V，#=5.000 V", "PASS"),
            ("断开翻转", "约 11.56 ms 跨过逻辑阈值", "PASS"),
        ],
        "conclusion": "最新网表已将 *_ACTUAL 接至缓冲器输出、*# 接至反相器输出，闭合/断开极性与接口表一致。",
        "image": "relay_detection_waveform.png",
        "caption": "图 1 继电器检测板瞬态波形",
    },
    {
        "title": "4. EBS 板",
        "purpose": "验证 EBS 工作输出与看门狗状态、EBS_TRIG 触发输入之间的逻辑关系。正常状态下 EBS_WORK 输出 24 V；看门狗失效或触发故障时输出拉低。",
        "timing": [
            ("0-2 ms", "看门狗正常，EBS_TRIG=0 V", "EBS_WORK=24 V，AS_LOCK 为约 5 V 逻辑高"),
            ("2.01-4 ms", "看门狗故障", "EBS_WORK=0 V"),
            ("5-7 ms", "EBS_TRIG=5 V 触发故障", "EBS_WORK=0 V"),
            ("7.01-9 ms", "恢复正常", "EBS_WORK 恢复 24 V"),
        ],
        "results": [
            ("1 ms 正常", "EBS_WORK=24 V，AS_LOCK=4.225 V", "PASS"),
            ("3 ms 看门狗故障", "EBS_WORK=0 V", "PASS"),
            ("6 ms 触发故障", "EBS_WORK=0 V", "PASS"),
            ("8 ms 恢复", "EBS_WORK=24 V", "PASS"),
        ],
        "conclusion": "EBS 输出在正常、看门狗故障、触发故障和恢复四个阶段均符合接口表描述。",
        "image": "ebs_waveform.png",
        "caption": "图 2 EBS 板瞬态波形",
    },
    {
        "title": "5. 锁存板",
        "purpose": "验证 IMD/BMS/EBS/RES 错误输入进入锁存后的安全回路行为。图中展示 IMD 通道代表时序：故障出现后 SDC_OUT 断开，并在故障输入恢复后继续保持，直到复位信号拉低。",
        "timing": [
            ("0-2 ms", "IMD_Error=12 V，复位为高", "SDC_OUT=24 V"),
            ("2.01-3 ms", "IMD_Error 拉低为故障", "IMD_STATE 低，SDC_OUT=0 V"),
            ("3.01-5 ms", "IMD_Error 恢复正常", "锁存保持，SDC_OUT 仍为 0 V"),
            ("5.01-5.5 ms", "IMD_Reset=0 V", "清除锁存"),
            ("5.51 ms 起", "复位释放", "IMD_STATE 高，SDC_OUT 恢复 24 V"),
        ],
        "results": [
            ("1 ms 正常", "SDC_OUT=24 V", "PASS"),
            ("2.5 ms 故障", "IMD_STATE≈0 V，SDC_OUT=0 V", "PASS"),
            ("4 ms 保持", "故障输入恢复后仍锁存，SDC_OUT=0 V", "PASS"),
            ("6 ms 复位后", "IMD_STATE≈5 V，SDC_OUT=24 V", "PASS"),
        ],
        "conclusion": "锁存、保持和低电平复位动作均符合接口表描述。",
        "image": "lock_latch_waveform.png",
        "caption": "图 3 锁存板瞬态波形",
    },
    {
        "title": "6. BSPD 板",
        "purpose": "验证 BSPD 在电流超过 11.11 mA 且油压/制动条件成立后，小于 500 ms 切断安全回路，并在冲突消失后保持断开，直到 BSPD 电源复位。由于接口表未给出油压阈值，仿真将油压条件作为已满足的布尔输入。",
        "timing": [
            ("0-100 ms", "故障条件未成立", "SDC_OUT=24 V"),
            ("100.93 ms 起", "电流 > 11.11 mA 且油压条件成立", "开始 555 延时计时"),
            ("250 ms", "延时尚未结束", "SDC_OUT 仍保持 24 V"),
            ("415.22 ms", "延时触发", "SDC_OUT 由 24 V 切换为 0 V"),
            ("601 ms", "故障成立后约 500 ms", "SDC_OUT=0 V"),
            ("2.00 s", "冲突条件消失", "SDC_OUT 保持 0 V"),
            ("12.05 s", "冲突消失超过 10 s", "SDC_OUT 仍保持 0 V"),
            ("12.20-12.30 s", "BSPD 电源复位/复位信号拉低", "清除锁存"),
            ("12.45 s", "复位释放且冲突不存在", "SDC_OUT 恢复 24 V"),
        ],
        "results": [
            ("200 ms 故障条件", "FAULT_COND=5 V", "PASS"),
            ("250 ms 延时前", "SDC_OUT=24 V", "PASS"),
            ("断开延迟", "314.29 ms，小于 500 ms", "PASS"),
            ("故障+500 ms", "SDC_OUT=0 V", "PASS"),
            ("冲突消失+10 s", "SDC_OUT=0 V，仍保持断开", "PASS"),
            ("BSPD 复位后", "SDC_OUT=24 V", "PASS"),
        ],
        "conclusion": "在故障条件已满足的前提下，BSPD 于 314.29 ms 后断开，小于 500 ms；冲突消失超过 10 s 后未误闭合，BSPD 电源复位后恢复，满足本次时间要求。油压阈值仍需补充后才能做完整通过判定。",
        "image": "bspd_waveform.png",
        "caption": "图 4 BSPD 板断开延迟波形",
        "extra_figures": [
            ("bspd_hold_waveform.png", "图 5 BSPD 板保持断开与电源复位波形"),
        ],
    },
    {
        "title": "7. TSAL 板",
        "purpose": "验证 TSAL 红灯/绿灯阴极输出逻辑，以及红灯在高压存在时的连续闪烁频率和占空比。输出为低电平有效：阴极被拉低时对应灯亮，高电平时灯灭。",
        "timing": [
            ("0-1.20 s", "控制器/电池高压缩放电压大于 0.6 V，继电器未全断开", "红灯按 555 周期连续闪烁，GREEN_OUT_K=5 V"),
            ("50 ms / 250 ms", "高压存在时取红灯亮/灭两个采样点", "50 ms 红灯亮，250 ms 红灯灭"),
            ("360.36-720.72 ms", "取相邻两个红灯点亮上升沿", "测量一个完整闪烁周期"),
            ("1.201 s 起", "电压低于安全阈值且继电器状态为全断开", "RED_OUT_K=5 V，GREEN_OUT_K=0 V"),
            ("1.30 s", "安全状态稳定采样", "验证红绿灯输出极性"),
        ],
        "results": [
            ("50 ms 红灯亮", "RED_OUT_K=0 V，GREEN_OUT_K=5 V", "PASS"),
            ("250 ms 红灯灭", "RED_OUT_K=5 V，GREEN_OUT_K=5 V", "PASS"),
            ("闪烁周期", "360.36 ms", "PASS"),
            ("闪烁频率", "2.775 Hz，满足 2-5 Hz", "PASS"),
            ("占空比", "51.92%，接近 50%", "PASS"),
            ("1.30 s 安全", "RED_OUT_K=5 V，GREEN_OUT_K=0 V，V_SAFE=5 V，RELAYS_OPEN=5 V", "PASS"),
        ],
        "conclusion": "高压存在时红灯以 2.775 Hz、51.92% 占空比连续闪烁，满足 2-5 Hz 和 50% 占空比附近的要求；电压安全且继电器全断开时绿灯有效，符合接口表描述。",
        "image": "tsal_waveform.png",
        "caption": "图 6 TSAL 板红灯闪烁与安全状态波形",
    },
]

for board in boards:
    add_board_section(doc, board)

doc.save(OUT)
print(OUT)
